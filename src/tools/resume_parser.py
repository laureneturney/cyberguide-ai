"""Resume parsing — convert an uploaded PDF/DOCX/TXT into plain text.

Lazy-imports the heavy parsers (pypdf, python-docx) so a missing wheel only
breaks the file format that needs it; everything else degrades to text.
"""
from __future__ import annotations

import io
from dataclasses import dataclass


@dataclass
class ParsedResume:
    text: str
    source: str  # "pdf" | "docx" | "txt" | "paste"
    char_count: int
    notes: str = ""


def parse_upload(filename: str, data: bytes) -> ParsedResume:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _parse_pdf(data)
    if name.endswith(".docx"):
        return _parse_docx(data)
    # Treat anything else as plain text — works for .txt, .md, copy/paste.
    text = _decode_bytes(data)
    return ParsedResume(text=text.strip(), source="txt", char_count=len(text))


def parse_paste(text: str) -> ParsedResume:
    text = (text or "").strip()
    return ParsedResume(text=text, source="paste", char_count=len(text))


def _parse_pdf(data: bytes) -> ParsedResume:
    try:
        from pypdf import PdfReader  # lazy
    except Exception as exc:  # noqa: BLE001
        return ParsedResume(text="", source="pdf", char_count=0,
                            notes=f"pypdf unavailable: {exc}")
    try:
        reader = PdfReader(io.BytesIO(data))
        chunks = []
        for page in reader.pages:
            try:
                chunks.append(page.extract_text() or "")
            except Exception:
                continue
        text = _normalize("\n".join(chunks))
        return ParsedResume(text=text, source="pdf", char_count=len(text))
    except Exception as exc:  # noqa: BLE001
        return ParsedResume(text="", source="pdf", char_count=0,
                            notes=f"PDF parse failed: {exc}")


def _parse_docx(data: bytes) -> ParsedResume:
    try:
        from docx import Document  # lazy
    except Exception as exc:  # noqa: BLE001
        return ParsedResume(text="", source="docx", char_count=0,
                            notes=f"python-docx unavailable: {exc}")
    try:
        doc = Document(io.BytesIO(data))
        paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        # Pull table cells too — résumés often store skills/timelines in tables.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        paras.append(cell.text.strip())
        text = _normalize("\n".join(paras))
        return ParsedResume(text=text, source="docx", char_count=len(text))
    except Exception as exc:  # noqa: BLE001
        return ParsedResume(text="", source="docx", char_count=0,
                            notes=f"DOCX parse failed: {exc}")


def _decode_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _normalize(text: str) -> str:
    # Collapse runs of spaces, keep line structure (helpful for the LLM).
    out_lines = []
    for line in text.splitlines():
        line = " ".join(line.split())
        if line:
            out_lines.append(line)
    return "\n".join(out_lines).strip()
